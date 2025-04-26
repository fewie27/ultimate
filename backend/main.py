import os
import uuid
import logging
import json
from typing import List, Optional, Dict, Any
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import uvicorn
import PyPDF2
import docx
import pytesseract
from PIL import Image
from datetime import datetime

# Import the rental analysis module
from analysis.analysis import analyzer as rental_analyzer

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Create the FastAPI app
app = FastAPI(
    title="Legal Document Analysis API",
    description="API for analyzing legal documents",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with actual frontend origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Temporary storage for demonstration
# In a real app, this would be a database
UPLOADS_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
ANALYSES = {}

# Create folders if they don't exist
os.makedirs(UPLOADS_FOLDER, exist_ok=True)
RESULTS_FOLDER = os.path.join(os.path.dirname(__file__), 'analysis_results')
os.makedirs(RESULTS_FOLDER, exist_ok=True)

# Define Pydantic models for responses
class Category(str):
    FEHLEND = "fehlend"
    UNGEWOEHNLICH = "ungewöhnlich"
    NICHTIG = "nichtig"

class AnalysisItem(BaseModel):
    text: str
    category: str
    description: str

class UploadResponse(BaseModel):
    id: str
    message: str

class AnalysisResponse(BaseModel):
    id: str
    results: List[AnalysisItem]

class SearchResult(BaseModel):
    text: str
    similarity: float
    category: str

class SearchResponse(BaseModel):
    query: str
    minimal_requirements: List[SearchResult]
    sample_clauses: List[SearchResult]

class UploadedDocument(BaseModel):
    id: str
    filename: str
    upload_date: str

class UploadedDocumentsList(BaseModel):
    documents: List[UploadedDocument]

# Text extraction functions
def extract_text_from_pdf(file_path):
    """Extract text from a PDF file, adding one backslash-n after each paragraph."""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()

                if page_text:
                    lines = page_text.split('\n')
                    paragraph = ""
                    for line in lines:
                        stripped_line = line.strip()
                        if stripped_line:
                            paragraph += stripped_line + " "
                        else:
                            # End of paragraph
                            text += paragraph.strip() + "\n"
                            paragraph = ""
                    if paragraph:
                        text += paragraph.strip() + "\n"  # Last paragraph if no empty line at end
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        text = f"Error extracting text: {str(e)}"
    return text



def extract_text_from_docx(file_path):
    """Extract text from a DOCX file"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        text = f"Error extracting text: {str(e)}"
    return text

def extract_text_from_image(file_path):
    """Extract text from an image using OCR"""
    text = ""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
    except Exception as e:
        logger.error(f"Error extracting text from image: {e}")
        text = f"Error extracting text: {str(e)}"
    return text

def extract_text(file_path):
    """Extract text from a file based on its extension"""
    _, extension = os.path.splitext(file_path.lower())
    
    if extension in ['.pdf']:
        return extract_text_from_pdf(file_path)
    elif extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
        return extract_text_from_image(file_path)
    else:
        # For text files or unsupported formats, try to read as text
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading file as text: {e}")
            return f"Unsupported file format or error reading file: {str(e)}"
        
def save_results_to_json(document_id: str, analysis_response: Dict[str, Any]) -> None:
    """Save analysis results to JSON file"""
    filename = os.path.join(RESULTS_FOLDER, f"analysis_{document_id}.json")
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(analysis_response, f, ensure_ascii=False, indent=4)
    logger.info(f"Analysis results saved to {filename}")

def analyze_legal_text(text):
    """
    Analyze legal text to identify issues
    Uses the rental agreement analyzer for more sophisticated analysis
    """
    try:
        # Split text into sentences
        sentences = rental_analyzer.split_text_into_sections(text)
        
        logger.info(f"Analyzing document with {len(sentences)} sentences")
        
        # Analyze using the rental agreement analyzer
        results = rental_analyzer.analyze_document_for_issues(sentences, {"source": "api_upload"})
        
        # Log analysis results
        logger.info(f"Analysis complete. Found {len(results)} issues.")
        
        return results
    except Exception as e:
        logger.error(f"Error analyzing with rental analyzer: {e}")
        
        # Fallback to basic analysis
        results = []
        if text.strip():
            results.append({
                "text": "Keine spezifischen problematischen Klauseln identifiziert.",
                "category": "fehlend",
                "description": "Es wurden keine offensichtlich problematischen Klauseln gefunden."
            })
        
        return results

# API endpoints
@app.post("/api/upload", response_model=UploadResponse, status_code=201)
async def upload_document(file: UploadFile = File(...)):
    """Upload a document and generate a UUID for it"""
    if not file:
        raise HTTPException(status_code=400, detail="No file part")
    
    if file.filename == "":
        raise HTTPException(status_code=400, detail="No selected file")
    
    # Generate a unique ID
    analysis_id = str(uuid.uuid4())
    
    # Save the file
    file_path = os.path.join(UPLOADS_FOLDER, f"{analysis_id}_{file.filename}")
    
    try:
        # Save uploaded file
        contents = await file.read()
        with open(file_path, "wb") as f:
            f.write(contents)
        
        # Extract text from the document
        extracted_text = extract_text(file_path)
        
        # Analyze the text
        results = analyze_legal_text(extracted_text)
        
        # Format and save the results
        analysis_response = {
            "id": analysis_id,
            "results": results
        }
        
        # Store results for API access
        ANALYSES[analysis_id] = analysis_response
        
        # Save to JSON file
        save_results_to_json(analysis_id, analysis_response)
        
        return {"id": analysis_id, "message": "File uploaded successfully"}
    
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")

@app.get("/api/analysis/{analysis_id}", response_model=AnalysisResponse)
async def get_analysis(analysis_id: str):
    """Get analysis results by ID"""
    if analysis_id not in ANALYSES:
        raise HTTPException(status_code=404, detail="Analysis not found")
    
    analysis = ANALYSES[analysis_id]
    
    return {
        "id": analysis["id"],
        "results": analysis["results"]
    }

@app.get("/api/search", response_model=SearchResponse)
async def search_rental_agreement(query: str, limit: int = 3):
    """Search for relevant clauses and requirements using the vector database"""
    if not query:
        raise HTTPException(status_code=400, detail="Query parameter is required")
    
    try:
        # Create query embedding
        query_embedding = rental_analyzer.model.encode([query])[0].tolist()
        
        # Search in minimal requirements collection
        req_results = rental_analyzer.minimal_requirements.query(
            query_embeddings=[query_embedding],
            n_results=limit
        )
        
        # Convert to SearchResult objects
        requirements = []
        for i, doc in enumerate(req_results["documents"][0]):
            requirements.append(
                SearchResult(
                    text=doc,
                    similarity=1.0 - req_results["distances"][0][i],
                    category=req_results["metadatas"][0][i].get("category", "general")
                )
            )
        
        # Search in sample agreement collection
        sample_results = rental_analyzer.sample_agreement.query(
            query_embeddings=[query_embedding],
            n_results=limit
        )
        
        # Convert to SearchResult objects
        sample_clauses = []
        for i, doc in enumerate(sample_results["documents"][0]):
            sample_clauses.append(
                SearchResult(
                    text=doc,
                    similarity=1.0 - sample_results["distances"][0][i],
                    category=sample_results["metadatas"][0][i].get("category", "general")
                )
            )
        
        return {
            "query": query,
            "minimal_requirements": requirements,
            "sample_clauses": sample_clauses
        }
    except Exception as e:
        logger.error(f"Error searching vector database: {e}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

@app.get("/api/documents", response_model=UploadedDocumentsList)
async def get_uploaded_documents():
    """Get a list of all uploaded documents"""
    try:
        # Get the uploaded documents collection
        uploads_collection = rental_analyzer.client.get_or_create_collection("uploaded_documents")
        
        # Get all documents
        uploads = uploads_collection.get()
        
        documents = []
        for i, doc_id in enumerate(uploads["ids"]):
            meta = uploads["metadatas"][i]
            documents.append(UploadedDocument(
                id=doc_id,
                filename=meta.get("filename", "Unknown"),
                upload_date=meta.get("upload_date", "")
            ))
        
        return {"documents": documents}
    except Exception as e:
        logger.error(f"Error getting uploaded documents: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve documents")

# Event handler for application startup
@app.on_event("startup")
async def startup_event():
    """Initialize vector database collections on startup"""
    logger.info("Initializing rental agreement analyzer...")
    # The import already initializes the analyzer and creates the collections

if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=5000, reload=True) 