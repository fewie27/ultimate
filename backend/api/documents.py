import os
import uuid
import werkzeug
from flask import jsonify, request
import PyPDF2
import docx
import pytesseract
from PIL import Image
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Temporary storage for demonstration
# In a real app, this would be a database
UPLOADS_FOLDER = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads')
ANALYSES = {}

# Create uploads folder if it doesn't exist
os.makedirs(UPLOADS_FOLDER, exist_ok=True)

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        text = f"Error extracting text: {str(e)}"
    return text

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        text = f"Error extracting text: {str(e)}"
    return text

def extract_text_from_image(file_path):
    """Extract text from an image using OCR"""
    text = ""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
    except Exception as e:
        logger.error(f"Error extracting text from image: {e}")
        text = f"Error extracting text: {str(e)}"
    return text

def extract_text(file_path):
    """Extract text from a file based on its extension"""
    _, extension = os.path.splitext(file_path.lower())
    
    if extension in ['.pdf']:
        return extract_text_from_pdf(file_path)
    elif extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
        return extract_text_from_image(file_path)
    else:
        # For text files or unsupported formats, try to read as text
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading file as text: {e}")
            return f"Unsupported file format or error reading file: {str(e)}"

def analyze_legal_text(text):
    """
    Analyze legal text to identify issues
    This is a placeholder for actual analysis logic
    In a real application, this would use NLP/ML techniques
    """
    results = []
    
    # Simple pattern matching for demonstration purposes
    patterns = [
        {
            "keyword": "automatisch",
            "category": "ungewöhnlich",
            "description": "Automatische Verlängerungsklausel"
        },
        {
            "keyword": "verlängert sich",
            "category": "ungewöhnlich",
            "description": "Automatische Verlängerungsklausel"
        },
        {
            "keyword": "Kündigungsfrist",
            "category": "ungewöhnlich",
            "description": "Möglicherweise lange Kündigungsfrist"
        },
        {
            "keyword": "Monate vor",
            "category": "ungewöhnlich",
            "description": "Lange Kündigungsfrist"
        },
        {
            "keyword": "Gerichtsstand",
            "category": "nichtig",
            "description": "Potenziell unzulässige Gerichtsstandsklausel"
        },
        {
            "keyword": "ausschließlich",
            "category": "nichtig",
            "description": "Potenziell unzulässige Ausschlussklausel"
        }
    ]
    
    # Split text into paragraphs for analysis
    paragraphs = text.split('\n')
    for paragraph in paragraphs:
        if not paragraph.strip():
            continue
            
        for pattern in patterns:
            if pattern["keyword"].lower() in paragraph.lower():
                results.append({
                    "text": paragraph.strip(),
                    "category": pattern["category"],
                    "description": pattern["description"]
                })
                break
    
    # Default response if no patterns match
    if not results and text.strip():
        results.append({
            "text": "Keine spezifischen problematischen Klauseln identifiziert.",
            "category": "fehlend",
            "description": "Es wurden keine offensichtlich problematischen Klauseln gefunden."
        })
    
    return results

def uploadDocument():
    """Upload a document and generate a UUID for it"""
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if file:
        # Generate a unique ID
        analysis_id = str(uuid.uuid4())
        
        # Save the file
        filename = werkzeug.utils.secure_filename(file.filename)
        file_path = os.path.join(UPLOADS_FOLDER, f"{analysis_id}_{filename}")
        file.save(file_path)
        
        # Extract text from the document
        extracted_text = extract_text(file_path)
        
        # Analyze the text
        results = analyze_legal_text(extracted_text)
        
        # Store the analysis results
        ANALYSES[analysis_id] = {
            "id": analysis_id,
            "file_path": file_path,
            "text": extracted_text,
            "results": results
        }
        
        return jsonify({"id": analysis_id, "message": "File uploaded successfully"}), 201
    
    return jsonify({"error": "File upload failed"}), 400

def getAnalysisById(analysisId):
    """Get analysis results by ID"""
    if analysisId not in ANALYSES:
        return jsonify({"error": "Analysis not found"}), 404
    
    analysis = ANALYSES[analysisId]
    
    return jsonify({
        "id": analysis["id"],
        "results": analysis["results"]
    }), 200 